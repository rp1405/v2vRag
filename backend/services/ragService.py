from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import Chroma
from langchain.embeddings.base import Embeddings
from langchain.chat_models.base import BaseChatModel
from prompts.rag_system_prompt import rag_system_prompt
from PyPDF2 import PdfReader
import io
import chardet
import os
from typing import Generator, AsyncGenerator, Optional, List, Tuple, Union, BinaryIO
from services.elevenLabsService import ElevenLabsService
from services.embeddingsFactory import EmbeddingsFactory
from services.llmFactory import LLMFactory
from contextlib import contextmanager
from prompts.rag_system_prompt import SystemPrompt
from loguru import logger
import time
from services.namedEntityService import NamedEntityService
from datetime import datetime
from docx import Document  # Add this import for DOCX support
import pdfplumber


class DocumentContext:
    def __init__(self, role: str, content: str, metadata: dict):
        self.role = role
        self.content = content
        self.metadata = metadata


class RAGService:
    def __init__(
        self,
        embeddings: str,
        llm: str,
        chunk_size: int = 50,
        query_language: str = "english",
    ):
        self.vector_store: Optional[Chroma] = None
        self.embeddings = EmbeddingsFactory(embeddings).get_embeddings()
        self.llm = LLMFactory(llm).get_llm()
        self.context: List[DocumentContext] = []
        self.system_prompt = SystemPrompt(llm)
        self.user_prompt: Optional[str] = None
        self.eleven_labs = ElevenLabsService()
        self.chunk_size = chunk_size
        self.model_name = llm
        self.ner_service = NamedEntityService()
        self.chunks = []
        self.query_language = query_language

    def _fix_leading_punctuation(self, chunks):
        fixed_chunks = []
        for chunk in chunks:
            if fixed_chunks and chunk and chunk[0] in ".!?":
                # Move punctuation to end of previous chunk
                fixed_chunks[-1] += chunk[0]
                chunk = chunk[1:]
            chunk = chunk.strip()
            fixed_chunks.append(chunk)
        return fixed_chunks

    @contextmanager
    def _safe_file_handle(self, file_content: bytes) -> Generator[BinaryIO, None, None]:
        """Safely handle file operations with proper cleanup."""
        file_handle = io.BytesIO(file_content)
        try:
            yield file_handle
        finally:
            file_handle.close()

    def format_table_as_markdown(self, table, headers=None):
        """Format a table as markdown string"""
        if not table:
            return ""

        # Replace None with empty strings
        cleaned_table = [
            [cell.strip() if cell else "" for cell in row] for row in table
        ]

        # Use first row as header if not provided
        if not headers:
            headers = cleaned_table[0]
            rows = cleaned_table[1:]
        else:
            rows = cleaned_table

        # Normalize multi-line headers and cells
        headers = [" ".join(col.split()) for col in headers]
        rows = [[" ".join(cell.split()) for cell in row] for row in rows]

        # Build Markdown table
        markdown = "| " + " | ".join(headers) + " |\n"
        markdown += "| " + " | ".join(["---"] * len(headers)) + " |\n"
        for row in rows:
            # pad row to match header length
            padded_row = row + [""] * (len(headers) - len(row))
            markdown += "| " + " | ".join(padded_row) + " |\n"
        return markdown

    def extract_pdf_text_with_tables(self, pdf_content):
        """Extract text and tables from PDF content using pdfplumber"""
        try:
            pdf_file = io.BytesIO(pdf_content)
            processed_tables = set()  # Keep track of processed tables
            all_content = []

            with pdfplumber.open(pdf_file) as pdf:
                for page in pdf.pages:
                    page_content = []

                    # Extract tables first
                    tables = page.extract_tables()
                    for table in tables:
                        if not table:  # Skip empty tables
                            continue

                        # Create a string representation of the table for deduplication
                        table_str = str(table)
                        if table_str not in processed_tables:
                            processed_tables.add(table_str)
                            markdown_table = self.format_table_as_markdown(table)
                            if markdown_table.strip():  # Only add non-empty tables
                                page_content.append(("table", markdown_table))

                    # Extract text after tables
                    page_text = page.extract_text() or ""
                    if page_text.strip():
                        page_content.append(("text", page_text))

                    # Sort content to maintain original order (tables followed by text)
                    page_content.sort(key=lambda x: x[0] != "table")

                    # Add page content to main content list
                    all_content.extend([content for _, content in page_content])

            # Join all content with appropriate spacing
            final_text = "\n".join(all_content)
            logger.info("Successfully extracted text and tables from PDF")
            return final_text.strip()
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {e}")
            return ""

    def extract_text_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF content with proper error handling."""
        try:
            with self._safe_file_handle(file_content) as pdf_file:
                pdf_reader = PdfReader(pdf_file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except Exception as e:
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")

    def extract_text_from_docx(self, file_content: bytes) -> str:
        """Extract text and tables from DOCX content with proper error handling."""
        try:
            with self._safe_file_handle(file_content) as docx_file:
                doc = Document(docx_file)
                all_content = []

                # Process paragraphs and tables in order
                for element in doc.element.body:
                    if element.tag.endswith("p"):  # Paragraph
                        paragraph = doc.paragraphs[len(all_content)]
                        text = paragraph.text.strip()
                        if text:
                            all_content.append(("text", text))
                    elif element.tag.endswith("tbl"):  # Table
                        # Find the table index
                        table_index = (
                            len(
                                [
                                    t
                                    for t in doc.tables
                                    if t._element is element
                                    or t._element in element.iter()
                                ]
                            )
                            - 1
                        )
                        if table_index >= 0:
                            table = doc.tables[table_index]
                            if table.rows:
                                # Extract headers from first row
                                headers = [
                                    cell.text.strip() for cell in table.rows[0].cells
                                ]
                                # Extract data rows
                                rows = []
                                for row in table.rows[1:]:
                                    row_data = [cell.text.strip() for cell in row.cells]
                                    if any(row_data):  # Only add non-empty rows
                                        rows.append(row_data)

                                if (
                                    headers and rows
                                ):  # Only process if we have both headers and data
                                    markdown_table = self.format_table_as_markdown(
                                        rows, headers
                                    )
                                    if markdown_table.strip():
                                        all_content.append(("table", markdown_table))

                # Join all content with appropriate spacing
                final_text = []
                for content_type, content in all_content:
                    if content_type == "table":
                        final_text.append(
                            "\n" + content + "\n"
                        )  # Add extra newlines around tables
                    else:
                        final_text.append(content)

                return "\n".join(final_text).strip()
        except Exception as e:
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")

    def process_file(
        self, file_content: bytes, filename: str, user_prompt: str
    ) -> None:
        """Process a file and create vector store with proper error handling."""
        try:
            # Clean up any existing vector store
            if self.vector_store:
                self.vector_store = None

            if filename.lower().endswith(".pdf"):
                content = self.extract_pdf_text_with_tables(file_content)
                if not content.strip():
                    # Fallback to basic extraction if table-aware extraction fails
                    content = self.extract_text_from_pdf(file_content)
            elif filename.lower().endswith(".docx"):
                content = self.extract_text_from_docx(file_content)
                if not content.strip():
                    # Fallback to basic extraction if table-aware extraction fails
                    content = self.extract_text_from_pdf(file_content)
            else:
                encoding = chardet.detect(file_content)["encoding"]
                if not encoding:
                    encoding = "utf-8"
                try:
                    content = file_content.decode(encoding)
                except UnicodeDecodeError:
                    content = file_content.decode("latin-1")
            self.user_prompt = user_prompt
            self.generate_vector_store(content, filename)
        except Exception as e:
            raise ValueError(f"Failed to process file {filename}: {str(e)}")

    def generate_vector_store(self, content: str, filename: str = "default") -> None:
        try:
            # Adjust chunk size and overlap for better table handling
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=1000,
                chunk_overlap=100,
                length_function=len,
                separators=[
                    "\n\n",
                    "\n",
                    ".",
                    "?",
                    "!",
                    "|",
                ],  # Added | as separator for tables
            )
            chunks = text_splitter.split_text(content)
            chunks = self._fix_leading_punctuation(chunks)
            # Filter out empty chunks and normalize whitespace
            chunks = [chunk.strip() for chunk in chunks if chunk.strip()]
            self.chunks = chunks
            logger.info(f"Using embedding: {self.embeddings}")
            # Create a new vector store with a unique collection name
            collection_name = f"doc_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
            self.vector_store = Chroma.from_texts(
                chunks,
                self.embeddings,
                metadatas=[
                    {
                        "source": filename,
                        "created_at": datetime.now().strftime("%Y%m%d_%H%M%S"),
                    }
                    for _ in chunks
                ],
                collection_name=collection_name,
            )
        except Exception as e:
            raise ValueError(f"Failed to generate vector store: {str(e)}")

    def generate_query(self, query: str) -> str:
        """Generate a query with proper error handling."""
        if not self.vector_store:
            raise ValueError(
                "Vector store not initialized. Please process a file first."
            )

        try:
            # Get relevant documents
            similarity_docs = [
                doc.page_content
                for doc in self.vector_store.similarity_search(query, k=4)
            ]

            keyword_docs, keyword_scores = self.ner_service.search_chunks_for_entities(
                query, self.chunks, k=4
            )

            context = similarity_docs + keyword_docs

            prev_context_text = ""
            for document in self.context:
                prev_context_text += (
                    f"{document.role.capitalize()}: {document.content}\n\n"
                )

            prompt = self.system_prompt.get_prompt(
                query,
                prev_context_text,
                context,
                self.user_prompt,
                style="fid",
                query_language=self.query_language,
            )
            logger.info(f"Prompt: {prompt}")
            return prompt
        except Exception as e:
            raise ValueError(f"Failed to generate query: {str(e)}")

    async def query_document_stream_async(
        self, query: str
    ) -> AsyncGenerator[Tuple[str, bytes], None]:
        """
        Stream the response from the LLM asynchronously with proper error handling.
        Returns a tuple of (text_chunk, audio_chunk)
        """
        if not self.llm:
            raise ValueError("LLM not initialized. Please check your configuration.")

        try:
            full_prompt = self.generate_query(query)
            self.context.append(
                DocumentContext(role="user", content=query, metadata={})
            )

            response_chunks = []
            current_buffer = ""
            current_len = 0

            async for chunk in self.llm.stream(full_prompt):
                if not hasattr(chunk, "content"):
                    continue

                current_buffer += chunk.content
                current_len += 1

                if current_len >= self.chunk_size:
                    try:
                        # Get audio for the current buffer
                        audio_chunks = []
                        for audio_chunk in self.eleven_labs.stream_text_to_speech(
                            current_buffer
                        ):
                            audio_chunks.append(audio_chunk)

                        # Yield both text and audio
                        yield current_buffer, b"".join(audio_chunks)

                        response_chunks.append(current_buffer)
                        current_buffer = ""
                        current_len = 0
                    except Exception as e:
                        raise ValueError(f"Failed to generate audio: {str(e)}")

            # Handle any remaining text
            if current_buffer:
                try:
                    audio_chunks = []
                    for audio_chunk in self.eleven_labs.stream_text_to_speech(
                        current_buffer
                    ):
                        audio_chunks.append(audio_chunk)

                    yield current_buffer, b"".join(audio_chunks)
                    response_chunks.append(current_buffer)
                except Exception as e:
                    raise ValueError(
                        f"Failed to generate audio for remaining text: {str(e)}"
                    )

            complete_response = "".join(response_chunks)
            self.context.append(
                DocumentContext(
                    role="assistant", content=complete_response, metadata={}
                )
            )
        except Exception as e:
            raise ValueError(f"Failed to process query: {str(e)}")

    def cleanup(self) -> None:
        """Clean up resources."""
        if self.vector_store:
            self.vector_store = None
        self.context.clear()
        self.user_prompt = None

    async def query_document(self, query: str) -> str:
        """
        Get the complete response for a query (non-streaming version).
        """
        if not self.llm:
            raise ValueError("Please upload a document first")

        full_prompt = self.generate_query(query)
        self.context.append(DocumentContext(role="user", content=query, metadata={}))

        response = await self.llm.get_complete_response(full_prompt)

        self.context.append(
            DocumentContext(role="assistant", content=response, metadata={})
        )

        return response
